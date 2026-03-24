"""
Document Processor — Phase 9
Extracts text from PDF, Excel/CSV, DOCX, and Email files.
All processing is in-memory — files are never persisted to disk.

Supported formats:
  - PDF (.pdf)         — pdfplumber
  - Excel (.xlsx/.xls) — openpyxl
  - CSV (.csv)         — stdlib csv
  - Word (.docx)       — python-docx
  - Email (.eml/.msg)  — stdlib email
"""
import io
import csv
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

# ── Max chars per document (to prevent token overflow) ───────────────────────
MAX_CHARS = 80_000  # ~20K tokens, enough for GPT-4 context


def extract_text(file_bytes: bytes, filename: str, content_type: str) -> Tuple[str, str]:
    """
    Extract plain text from a document.

    Returns:
        (text, detected_format) — text is the extracted content,
        detected_format is one of: pdf, excel, csv, docx, email, unknown
    """
    fname = filename.lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if fname.endswith(".pdf") or content_type == "application/pdf":
        return _extract_pdf(file_bytes), "pdf"

    # ── Excel ─────────────────────────────────────────────────────────────────
    if fname.endswith((".xlsx", ".xls")) or content_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_excel(file_bytes), "excel"

    # ── CSV ───────────────────────────────────────────────────────────────────
    if fname.endswith(".csv") or content_type == "text/csv":
        return _extract_csv(file_bytes), "csv"

    # ── Word DOCX ─────────────────────────────────────────────────────────────
    if fname.endswith(".docx") or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes), "docx"

    # ── Email ─────────────────────────────────────────────────────────────────
    if fname.endswith((".eml", ".msg")) or content_type in (
        "message/rfc822", "application/octet-stream"
    ):
        return _extract_email(file_bytes), "email"

    # ── Plain text fallback ───────────────────────────────────────────────────
    if content_type.startswith("text/"):
        text = file_bytes.decode("utf-8", errors="replace")
        return text[:MAX_CHARS], "text"

    raise ValueError(f"Unsupported file type: {filename} ({content_type})")


# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i+1}]\n{text.strip()}")
            return "\n\n".join(pages)[:MAX_CHARS]
    except ImportError:
        raise ImportError("pdfplumber not installed. Add to requirements.txt: pdfplumber")
    except Exception as e:
        logger.error(f"[DocProcessor] PDF extraction failed: {e}")
        raise ValueError(f"Could not extract text from PDF: {e}")


def _extract_excel(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
            if rows:
                parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        return "\n\n".join(parts)[:MAX_CHARS]
    except ImportError:
        raise ImportError("openpyxl not installed. Add to requirements.txt: openpyxl")
    except Exception as e:
        logger.error(f"[DocProcessor] Excel extraction failed: {e}")
        raise ValueError(f"Could not extract text from Excel: {e}")


def _extract_csv(data: bytes) -> str:
    try:
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = ["\t".join(row) for row in reader if any(c.strip() for c in row)]
        return "\n".join(rows)[:MAX_CHARS]
    except Exception as e:
        logger.error(f"[DocProcessor] CSV extraction failed: {e}")
        raise ValueError(f"Could not extract text from CSV: {e}")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                tables.append("\n".join(rows))

        parts = paragraphs
        if tables:
            parts += ["[Tables]"] + tables

        return "\n\n".join(parts)[:MAX_CHARS]
    except ImportError:
        raise ImportError("python-docx not installed. Add to requirements.txt: python-docx")
    except Exception as e:
        logger.error(f"[DocProcessor] DOCX extraction failed: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")


def _extract_email(data: bytes) -> str:
    try:
        import email as email_lib
        from email import policy as email_policy

        msg = email_lib.message_from_bytes(data, policy=email_policy.default)

        parts = []

        # Headers
        subject = msg.get("Subject", "")
        sender = msg.get("From", "")
        to = msg.get("To", "")
        date = msg.get("Date", "")

        if subject:
            parts.append(f"Subject: {subject}")
        if sender:
            parts.append(f"From: {sender}")
        if to:
            parts.append(f"To: {to}")
        if date:
            parts.append(f"Date: {date}")

        parts.append("")  # blank line before body

        # Body
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        parts.append(payload.decode(charset, errors="replace"))
                        break  # prefer plain text
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or "utf-8"
                parts.append(payload.decode(charset, errors="replace"))

        return "\n".join(parts)[:MAX_CHARS]
    except Exception as e:
        logger.error(f"[DocProcessor] Email extraction failed: {e}")
        raise ValueError(f"Could not extract text from email: {e}")
